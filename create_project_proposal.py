from docx import Document

proposal = Document()
proposal.add_heading('Healthcare Prediction System - Project Proposal', level=1)
proposal.add_paragraph('')

proposal.add_heading('1. Project Overview', level=2)
proposal.add_paragraph(
    'The Healthcare Prediction System is a Flask-based web application that enables users to register, log in, and receive disease predictions based on symptom input. ' 
    'The system stores user details and prediction history in MongoDB, while disease descriptions and precautions are shown together with the prediction result.'
)
proposal.add_paragraph(
    'The application includes a master admin dashboard that allows an administrator to view all registered users and their entire prediction history. ' 
    'Users and admins can also download a PDF report of each prediction.'
)

proposal.add_heading('2. Objectives', level=2)
proposal.add_paragraph('The primary objectives of this project are:')
proposal.add_paragraph('• Build a secure user registration and authentication system using Flask and MongoDB.', style='List Bullet')
proposal.add_paragraph('• Integrate a machine learning model to predict diseases from user-selected symptoms.', style='List Bullet')
proposal.add_paragraph('• Store prediction history and user profile data in MongoDB Atlas.', style='List Bullet')
proposal.add_paragraph('• Provide an admin interface for monitoring users and predictions.', style='List Bullet')
proposal.add_paragraph('• Generate downloadable PDF reports for prediction results.', style='List Bullet')

proposal.add_heading('3. Project Scope', level=2)
proposal.add_paragraph(
    'This project focuses on the following key components:'
)
proposal.add_paragraph('• User registration and login with password hashing.', style='List Bullet')
proposal.add_paragraph('• Symptom-based disease prediction using a pre-trained machine learning model.', style='List Bullet')
proposal.add_paragraph('• MongoDB storage for users and prediction records.', style='List Bullet')
proposal.add_paragraph('• Admin dashboard with user and history management.', style='List Bullet')
proposal.add_paragraph('• PDF report generation for each prediction.', style='List Bullet')
proposal.add_paragraph('• Multilingual support for symptom and disease names with English/Bengali labels.', style='List Bullet')

proposal.add_heading('4. System Architecture', level=2)
proposal.add_paragraph(
    'The system architecture includes the following major layers:'
)
proposal.add_paragraph('• Presentation layer: Flask templates and HTML forms for user interaction.', style='List Bullet')
proposal.add_paragraph('• Application layer: Flask routes and controllers in app.py.', style='List Bullet')
proposal.add_paragraph('• Data layer: MongoDB Atlas for persistent user and prediction data.', style='List Bullet')
proposal.add_paragraph('• Machine learning layer: trained RandomForest classifier and symptom-weighted feature engineering.', style='List Bullet')
proposal.add_paragraph('• Reporting layer: PDF generation using ReportLab.', style='List Bullet')

proposal.add_heading('5. Data and Model', level=2)
proposal.add_paragraph(
    'The machine learning component is trained using the provided disease dataset and symptom severity information. ' 
    'The training process creates the following serialized files:'
)
proposal.add_paragraph('• disease_model.pkl - trained RandomForest model', style='List Bullet')
proposal.add_paragraph('• label_encoder.pkl - label encoder for disease names', style='List Bullet')
proposal.add_paragraph('• symptom_weights.pkl - symptom severity weights', style='List Bullet')
proposal.add_paragraph('• all_symptoms.pkl - sorted symptom list for form dropdowns', style='List Bullet')
proposal.add_paragraph('• disease_info.pkl - disease descriptions and precautions', style='List Bullet')
proposal.add_paragraph(
    'The model predicts a disease from the selected symptoms, and the application displays the disease description, recommended precautions, and a localized Bengali name when available.'
)

proposal.add_heading('6. Features', level=2)
proposal.add_paragraph('Key features of the Healthcare Prediction System include:')
proposal.add_paragraph('• User registration and secure login with hashed passwords.', style='List Bullet')
proposal.add_paragraph('• Symptom selection and disease prediction flow.', style='List Bullet')
proposal.add_paragraph('• Admin account auto-creation based on environment variables.', style='List Bullet')
proposal.add_paragraph('• Prediction history storage and viewing for users and admin.', style='List Bullet')
proposal.add_paragraph('• PDF report download for prediction results.', style='List Bullet')
proposal.add_paragraph('• Bengali translations for symptoms and diseases.', style='List Bullet')

proposal.add_heading('7. Technology Stack', level=2)
proposal.add_paragraph('The project uses the following technologies:')
proposal.add_paragraph('• Python with Flask for back-end web development.', style='List Bullet')
proposal.add_paragraph('• MongoDB Atlas for cloud database storage.', style='List Bullet')
proposal.add_paragraph('• scikit-learn for machine learning model training and prediction.', style='List Bullet')
proposal.add_paragraph('• ReportLab for PDF report generation.', style='List Bullet')
proposal.add_paragraph('• python-dotenv for environment variable management.', style='List Bullet')
proposal.add_paragraph('• HTML/CSS templates for the front-end interface.', style='List Bullet')

proposal.add_heading('8. Deployment and Environment', level=2)
proposal.add_paragraph(
    'The application requires the following environment configuration before running:'
)
proposal.add_paragraph('• MONGO_URI set to the MongoDB Atlas connection string.', style='List Bullet')
proposal.add_paragraph('• ADMIN_EMAIL and ADMIN_PASSWORD for the master admin account.', style='List Bullet')
proposal.add_paragraph('• FLASK_SECRET_KEY for application session security.', style='List Bullet')
proposal.add_paragraph(
    'The project must also include the generated pickle files in the application folder and run on a system with the required Python packages installed.'
)

proposal.add_heading('9. Team Roles', level=2)
proposal.add_paragraph('Suggested responsibilities for team members:')
proposal.add_paragraph('• Front-end and templates: design and UX for login, dashboard, and prediction pages.', style='List Bullet')
proposal.add_paragraph('• Back-end and API: Flask route handling, authentication, and MongoDB integration.', style='List Bullet')
proposal.add_paragraph('• ML model and data preprocessing: training pipeline, symptom encoding, and prediction logic.', style='List Bullet')
proposal.add_paragraph('• Documentation and testing: proposal writing, user flow validation, and deployment testing.', style='List Bullet')

proposal.add_heading('10. Future Enhancements', level=2)
proposal.add_paragraph('Potential improvements for future development:')
proposal.add_paragraph('• Add stronger input validation and form error handling.', style='List Bullet')
proposal.add_paragraph('• Add role-based access control for multiple admin levels.', style='List Bullet')
proposal.add_paragraph('• Expand disease coverage and improve model accuracy with more data.', style='List Bullet')
proposal.add_paragraph('• Add a responsive front-end design for mobile devices.', style='List Bullet')
proposal.add_paragraph('• Add email notifications or alerts for users after prediction.', style='List Bullet')

proposal.add_heading('11. Conclusion', level=2)
proposal.add_paragraph(
    'The Healthcare Prediction System provides an end-to-end solution for symptom-based disease detection, ' 
    'user management, and reporting. It is a practical final-year project that combines web development, data science, and cloud database integration.'
)

proposal.save('Healthcare_Prediction_System_Project_Proposal.docx')
print('Proposal document created: Healthcare_Prediction_System_Project_Proposal.docx')
